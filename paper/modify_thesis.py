# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SRC = Path(r"D:\OJ-project\paper\thesis_work.docx")
OUT = Path(r"D:\OJ-project\paper\毕业设计-按意见修改版.docx")


doc = Document(str(SRC))

# Capture existing styles before changing paragraph order.
STYLE_H1 = doc.paragraphs[107].style
STYLE_H2 = doc.paragraphs[108].style
STYLE_H3 = doc.paragraphs[114].style
STYLE_BODY = doc.paragraphs[109].style
STYLE_CAPTION = doc.paragraphs[170].style
STYLE_ABS = doc.paragraphs[43].style


def paragraph(element, parent):
    return Paragraph(element, parent)


def clear(p, text, style=None):
    if style is not None:
        p.style = style
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    if text:
        p.add_run(text)
    return p


def after(p, text="", style=None):
    el = OxmlElement("w:p")
    p._p.addnext(el)
    new_p = paragraph(el, p._parent)
    if style is not None:
        new_p.style = style
    if text:
        new_p.add_run(text)
    return new_p


def delete(p):
    p._element.getparent().remove(p._element)


def cite(p, nums):
    for num in nums:
        run = p.add_run(f"[{num}]")
        run.font.superscript = True


def outline(p, level):
    p_pr = p._p.get_or_add_pPr()
    node = p_pr.find(qn("w:outlineLvl"))
    if node is None:
        node = OxmlElement("w:outlineLvl")
        p_pr.append(node)
    node.set(qn("w:val"), str(level))


def page_break_before(p):
    p_pr = p._p.get_or_add_pPr()
    if p_pr.find(qn("w:pageBreakBefore")) is None:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def add_toc_field(p):
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    for typ, val in [
        ("begin", None),
        ("instr", ' TOC \\o "1-3" \\h \\z \\u '),
        ("separate", None),
    ]:
        r = OxmlElement("w:r")
        if typ == "instr":
            child = OxmlElement("w:instrText")
            child.set(qn("xml:space"), "preserve")
            child.text = val
        else:
            child = OxmlElement("w:fldChar")
            child.set(qn("w:fldCharType"), typ)
        r.append(child)
        p._p.append(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录将在打开文档时自动更新"
    r.append(t)
    p._p.append(r)
    r = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)
    p._p.append(r)


settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

# 摘要，控制在 500 字以内。
abstract = (
    "针对程序设计课程中人工批改周期长、评测标准不统一、竞赛组织和学习过程数据难以沉淀等问题，本文设计并实现了一套基于 Spring Boot 与 Vue 的程序设计评测系统。"
    "系统采用前后端分离架构，后端集成 Spring Security、JWT、MyBatis-Plus、Flyway、MySQL 和 Redis，前端使用 Vue 3、Vite 与 Pinia 构建题库、代码工坊、赛事中枢、教师工作台和管理控制台等页面。"
    "系统实现了学生练习与竞赛、教师出题与数据分析、管理员配置与审计等功能，并在判题模块中设计本地执行与 Docker 沙箱执行双路径机制，结合隐藏测试用例、排行榜快照和封榜策略，提高代码评测的安全性、可维护性和教学适配性。"
    "测试结果表明，系统核心业务流程能够稳定运行，可为程序设计教学的信息化管理和竞赛训练提供实用支撑。"
)
clear(doc.paragraphs[43], abstract, STYLE_ABS)
delete(doc.paragraphs[44])
delete(doc.paragraphs[44])

# 1.3
p13 = next(p for p in doc.paragraphs if p.text.strip().startswith("1.3") and p.style.name != "Body Text")
clear(p13, "1.3  研究目标与内容", STYLE_H2)
outline(p13, 1)
old = []
collect = False
for p in list(doc.paragraphs):
    if p._element is p13._element:
        collect = True
        continue
    if collect and p.text.strip().startswith("1.4"):
        break
    if collect:
        old.append(p)
points = [
    "（1）完成面向三类角色的需求分析与功能划分。本文先从程序设计课程练习、竞赛训练和后台运维三个使用场景入手，梳理学生、教师、管理员的业务边界，再结合项目源码中的控制器、服务类和页面模块确定功能范围，使系统需求能够落到具体接口和页面上，避免停留在概念描述层面。",
    "（2）完成在线判题核心链路的设计与实现。系统围绕题目、测试用例、提交记录、判题结果和测试点结果建立数据关联，通过 JudgeServiceImpl 组织提交保存、测试数据读取、输出比对和状态回写，并支持 Java、C++、Python 三类语言评测，使学生提交代码后能够及时获得通过、答案错误、编译失败或运行超时等反馈。",
    "（3）完成 Docker 沙箱与本地判题双路径机制。针对在线执行代码存在的安全风险，系统在容器执行时设置禁网、只读文件系统、能力裁剪、内存、CPU 和进程数限制；当教学演示环境不具备容器条件时，又可按配置回退到本地判题，从而在安全性与可部署性之间取得平衡。",
    "（4）完成竞赛、讨论、教师分析和管理控制台等扩展功能。系统不只提供题目提交，还实现竞赛报名、实时榜单、快照榜单、帖子评论、私信交流、教学概览、操作日志、系统监控和论坛审核等功能，使平台能够覆盖课程训练、阶段竞赛和日常管理等多种教学场景。",
    "（5）完成系统测试与应用效果验证。本文依据后端单元测试、控制器测试、前端烟雾测试、竞赛端到端脚本和 k6 性能测试方案，对登录、题库、提交、竞赛、讨论、分析和后台管理等核心流程进行验证，结果表明系统主要功能能够按预期运行，并具有继续扩展到真实教学环境的应用价值。",
]
clear(old[0], points[0], STYLE_BODY)
cur = old[0]
for text in points[1:]:
    cur = after(cur, text, STYLE_BODY)
for p in old[1:]:
    delete(p)

# 2.1
p21 = next(p for p in doc.paragraphs if p.text.strip().startswith("2.1") and p.style.name != "Body Text")
seen = False
for p in doc.paragraphs:
    if seen and p.text.strip():
        clear(
            p,
            "本系统在 Windows 11 环境下开发和调试，后端运行环境为 JDK 17，构建工具为 Maven 3.8+，主要框架版本为 Spring Boot 2.7.18、Spring Security 5.x、MyBatis-Plus 3.5.3.1、Flyway 9.22.3，数据库采用 MySQL 8.0，缓存与限流状态存储采用 Redis 6.0，数据库驱动为 mysql-connector-j 8.0.33。前端运行环境采用 Node.js 18+ 与 npm，工程框架为 Vue 3.5.30、Vite 8.0.0、Pinia 3.0.4、Vue Router 5.0.3，并结合 Axios 1.13.6、CodeMirror 6 和 Tailwind CSS 3.4.17 完成页面交互与代码编辑体验。判题运行环境支持 javac/java、g++ 和 Python 解释器，容器化隔离依赖 Docker Desktop；测试工具包括 JUnit 5、Spring Boot Test、PowerShell 烟雾测试脚本和 k6 性能测试脚本。上述版本与项目依赖文件和运行说明保持一致，能够满足本系统开发、部署、调试和测试需要。",
            STYLE_BODY,
        )
        cite(p, [5, 8, 9, 10, 11, 12, 13])
        break
    if p._element is p21._element:
        seen = True

# 3.2
p32 = next(p for p in doc.paragraphs if p.text.strip().startswith("3.2") and p.style.name != "Body Text")
seen = False
for p in doc.paragraphs:
    if seen and p.text.strip():
        role_body = p
        break
    if p._element is p32._element:
        seen = True
clear(role_body, "学生角色主要关注学习训练链路是否顺畅。学生需要完成注册登录、题库浏览、题目详情查看、在线编写代码、提交评测、查看历史记录和测试点反馈等操作；在竞赛场景下，还需要报名参赛、查看实时榜单和赛后结果；在学习交流场景下，需要能够发帖、评论、关注用户和收发私信。系统通过题库中心、代码工坊、赛事中枢和讨论广场等页面把这些功能集中呈现，使学生能够围绕“读题—编码—提交—反馈—复盘”的流程持续练习。", STYLE_BODY)
cur = after(role_body, "教师角色主要关注教学组织和数据分析效率。教师需要创建和维护题目，配置样例与隐藏测试用例，组织课程练习或阶段竞赛，查看学生提交情况、语言分布、通过率和近期提交趋势，并在需要时导出统计数据。系统通过教师工作台、题目管理和竞赛管理接口支撑这些需求，使教师可以把更多精力放在题目设计、课堂反馈和个别辅导上，减少重复批改与手工统计工作。", STYLE_BODY)
after(cur, "管理员角色主要关注平台稳定、安全和可治理性。管理员需要维护用户账号和角色档案，调整系统配置，查看操作日志和运行监控，处理论坛内容审核，并观察判题结果和异常状态。系统通过管理控制台将用户管理、配置管理、日志审计、系统监控和内容治理放在统一入口中，使平台在多用户、多提交和多模块协同运行时能够保持清晰的权限边界和可追踪的管理记录。", STYLE_BODY)

# 3.3 + 3.4
p33 = next(p for p in doc.paragraphs if p.text.strip().startswith("3.3") and p.style.name != "Body Text")
clear(p33, "3.3  非功能性需求分析", STYLE_H2)
outline(p33, 1)
seen = False
for p in doc.paragraphs:
    if seen and p.text.strip():
        nf_body = p
        break
    if p._element is p33._element:
        seen = True
nfs = [
    "在安全性方面，系统需要对登录状态、角色权限、敏感接口和不可信代码执行进行重点控制。后端通过 Spring Security 与 JWT 识别用户身份，针对题目创建、竞赛管理、系统配置、论坛审核等操作进行角色校验；判题模块通过 Docker 沙箱限制网络、文件系统、进程数量、CPU 和内存资源，降低恶意代码或异常程序影响服务器的风险。",
    "在并发性方面，系统需要支撑同一时间内多名学生浏览题库、提交代码和查看竞赛榜单。后端在提交链路中记录提交状态并支持前端轮询查询，竞赛榜单通过成绩快照降低重复统计压力，Redis 可用于缓存热点状态和限流计数，从而缓解短时间高频提交对数据库和判题服务造成的压力。",
    "在响应时间方面，系统普通查询接口应尽量保持较短等待时间，题目列表、题目详情、竞赛列表和登录接口作为高频访问入口，需要配合分页查询、索引设计和缓存策略进行优化。性能测试方案以失败率低于 2%、95 分位响应时间低于 800ms 作为参考目标，保证系统在教学规模并发访问下仍具备较好的可用性。",
    "在可维护性方面，系统应保持前后端职责清楚、数据库结构可演进、测试脚本可重复执行。项目采用控制层、服务层和持久层分层组织后端代码，使用 Flyway 管理数据库迁移，并通过单元测试、控制器测试、烟雾测试和性能脚本对主要功能进行持续验证，为后续扩展新语言、新题型和分布式判题能力留下空间。",
]
clear(nf_body, nfs[0], STYLE_BODY)
cite(nf_body, [14])
cur = nf_body
for text in nfs[1:]:
    cur = after(cur, text, STYLE_BODY)
h34 = after(cur, "3.4  本章小结", STYLE_H2)
outline(h34, 1)
after(h34, "本章从可行性、角色需求和非功能性需求三个方面进行了分析。通过技术、经济和操作可行性的说明，以及对学生、教师、管理员三类用户需求的梳理，可以明确系统在教学场景中的核心业务要求；安全性、并发性、响应时间和可维护性等要求，则为后续总体架构设计、功能划分和数据库设计提供了约束依据。", STYLE_BODY)

for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("结合已有研究和实际平台建设情况"):
        cite(p, [1, 2])
    elif t.startswith("国外在线评测平台起步"):
        cite(p, [3, 4])
    elif t.startswith("后端使用 Spring Boot"):
        cite(p, [6, 7])
    elif t.startswith("项目在 onlinejudge-backend/scripts/perf"):
        cite(p, [15])

# 第 6 章测试表
pf = next(p for p in doc.paragraphs if p.text.strip().startswith("项目已经覆盖认证"))
caption = after(pf, "表6.1  核心功能测试用例与测试结果", STYLE_CAPTION)
table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["测试编号", "测试功能", "测试内容", "预期结果", "测试结果"]
for idx, head in enumerate(headers):
    table.rows[0].cells[idx].text = head
rows = [
    ["T01", "用户登录与权限", "使用学生、教师、管理员账号登录，并访问各自功能入口", "登录成功，越权访问被拦截", "通过"],
    ["T02", "题库浏览", "查询题目列表并进入题目详情页", "题目标题、难度、标签和题面内容正常显示", "通过"],
    ["T03", "测试用例维护", "教师新增、修改和删除题目测试用例", "测试用例保存成功，题目评测数据同步更新", "通过"],
    ["T04", "代码提交与判题", "学生提交 Java、C++ 或 Python 代码并轮询状态", "生成提交记录，返回 ACCEPTED、WRONG_ANSWER 等评测状态", "通过"],
    ["T05", "竞赛报名与排名", "创建竞赛、报名参赛并查询实时榜单和快照榜单", "榜单按通过题数和罚时排序，快照数据可查询", "通过"],
    ["T06", "讨论与社交", "发布帖子、评论、点赞、关注用户并发送私信", "讨论内容和社交关系正常保存，已读状态可更新", "通过"],
    ["T07", "教师分析", "访问教学概览并导出统计数据", "提交趋势、状态分布、语言分布和总体指标正常返回", "通过"],
    ["T08", "管理控制台", "管理员查看系统配置、操作日志、运行监控和论坛审核列表", "后台数据正常加载，敏感操作仅管理员可执行", "通过"],
    ["T09", "性能烟雾测试", "使用 k6 对登录、题目列表、题目详情和竞赛列表进行阶梯并发访问", "失败率低于 2%，p95 响应时间低于 800ms 的目标可作为验收标准", "符合预期"],
]
for row in rows:
    cells = table.add_row().cells
    for idx, val in enumerate(row):
        cells[idx].text = val
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for pp in cell.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in pp.runs:
                run.font.size = Pt(9)
caption._p.addnext(table._tbl)

# 结论
conclusion_head = next(p for p in doc.paragraphs if p.text.strip() == "结论")
seen = False
for p in doc.paragraphs:
    if seen and p.text.strip():
        conclusion_body = p
        break
    if p._element is conclusion_head._element:
        seen = True
conclusions = [
    "本文围绕“基于 Spring Boot 的程序设计评测系统的设计与实现”这一课题，完成了从需求分析、总体设计、数据库设计到详细实现和系统测试的完整研究。系统面向学生、教师和管理员三类角色，已经实现注册登录、题库浏览、在线提交、自动判题、竞赛报名、实时排名、讨论交流、教师分析、系统配置、日志审计和论坛审核等功能，能够覆盖程序设计课程练习和阶段竞赛中的主要业务流程。",
    "从系统实现效果看，项目采用 Spring Boot 与 Vue 前后端分离架构，后端通过 Spring Security、JWT、MyBatis-Plus、Flyway、MySQL 和 Redis 完成接口服务、权限控制和数据持久化，前端通过 Vue 3、Vite、Pinia 和 CodeMirror 提供题库、代码工坊、赛事中枢、教师工作台和管理控制台等交互页面。判题模块同时支持本地执行和 Docker 沙箱执行，并结合隐藏测试用例、测试点结果记录、时间限制补偿、竞赛快照和封榜机制，提高了系统在教学与竞赛场景中的适用性。",
    "从测试情况看，系统围绕登录权限、题库管理、测试用例维护、代码提交、竞赛排名、讨论社交、教师分析和后台管理等核心功能进行了验证，并结合烟雾测试、端到端测试和 k6 性能测试方案检查主要接口链路。测试结果表明，系统核心功能能够按预期运行，前后端接口交互较稳定，普通查询和竞赛相关操作能够满足课程训练场景下的基本使用要求。",
    "当然，系统仍存在进一步完善的空间。后续可在判题调度方面引入消息队列和分布式判题节点，以提升高并发提交下的处理能力；在安全治理方面加入更细粒度的内容审核、异常提交识别和敏感操作告警；在教学分析方面继续扩展班级、知识点和长期成长维度。随着这些功能逐步完善，系统可以更好地服务真实教学管理、程序设计训练和校内竞赛组织。",
]
clear(conclusion_body, conclusions[0], STYLE_BODY)
cur = conclusion_body
for text in conclusions[1:]:
    cur = after(cur, text, STYLE_BODY)

# 参考文献
ref_head = next(p for p in doc.paragraphs if p.text.strip() == "参考文献")
to_delete = []
collect = False
for p in list(doc.paragraphs):
    if p._element is ref_head._element:
        collect = True
        continue
    if collect and p.text.strip().replace(" ", "") == "致谢":
        break
    if collect and p.text.strip():
        to_delete.append(p)
for p in to_delete:
    delete(p)
refs = [
    "[1] 钟耀章,桂琼. ACM竞赛在线测评系统设计与实现[J]. 无线互联科技,2020(18):61-63.",
    "[2] 戴施伟,周凌珉,郑一泓. 基于SSM框架和RabbitMQ技术的OJ系统的设计与实现[J]. 计算机时代,2022(10):80-83.",
    "[3] Codeforces. Codeforces online programming contests and rating system[EB/OL]. https://codeforces.com/,2026-04-26.",
    "[4] DOMjudge Team. DOMjudge Manual 8.0[EB/OL]. https://www.domjudge.org/docs/manual/8.0/,2026-04-26.",
    "[5] Webb P, Syer D, Long J, et al. Spring Boot Reference Documentation 2.7.18[EB/OL]. https://docs.spring.io/spring-boot/docs/2.7.18/reference/htmlsingle/,2026-04-26.",
    "[6] Spring Security Team. Spring Security Reference 5.7[EB/OL]. https://docs.spring.io/spring-security/reference/,2026-04-26.",
    "[7] MyBatis-Plus Team. MyBatis-Plus Documentation 3.5.x[EB/OL]. https://baomidou.com/,2026-04-26.",
    "[8] Oracle Corporation. MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/,2026-04-26.",
    "[9] Redis Ltd. Redis Documentation 6.0[EB/OL]. https://redis.io/docs/,2026-04-26.",
    "[10] Docker Inc. Docker Engine Documentation: Resource constraints[EB/OL]. https://docs.docker.com/engine/containers/resource_constraints/,2026-04-26.",
    "[11] Vue.js Team. Vue.js Guide 3.x[EB/OL]. https://vuejs.org/guide/,2026-04-26.",
    "[12] Vite Team. Vite Guide[EB/OL]. https://vite.dev/guide/,2026-04-26.",
    "[13] Pinia Team. Pinia Documentation[EB/OL]. https://pinia.vuejs.org/,2026-04-26.",
    "[14] OWASP Foundation. OWASP Top 10:2021 Web Application Security Risks[EB/OL]. https://owasp.org/Top10/2021/,2026-04-26.",
    "[15] Grafana Labs. Grafana k6 Documentation[EB/OL]. https://grafana.com/docs/k6/latest/,2026-04-26.",
]
cur = ref_head
for ref in refs:
    cur = after(cur, ref, STYLE_BODY)

# Heading outline + page breaks.
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("第") and "章" in t and p.style.name in [STYLE_H1.name, "Heading 1"]:
        outline(p, 0)
        page_break_before(p)
    elif p.style.name in [STYLE_H2.name, "Heading 2"] and t[:1].isdigit():
        outline(p, 1)
    elif p.style.name in [STYLE_H3.name, "Heading 3"] and t[:1].isdigit():
        outline(p, 2)
    if t in ["结论", "参考文献", "致 谢", "致谢"]:
        outline(p, 0)
        page_break_before(p)

# Auto TOC field.
toc_title = None
for p in doc.paragraphs:
    if p.text.strip().replace(" ", "").replace("\u3000", "") == "目录":
        toc_title = p
        break
first_chapter = next(
    p for p in doc.paragraphs if p.text.strip().startswith("第 1 章") and p.style.name != "Body Text"
)
if toc_title is not None:
    stale = []
    collect = False
    for p in list(doc.paragraphs):
        if p._element is toc_title._element:
            collect = True
            continue
        if p._element is first_chapter._element:
            break
        if collect:
            stale.append(p)
    for p in stale:
        delete(p)
    add_toc_field(after(toc_title, style=STYLE_ABS))

doc.save(str(OUT))
print(OUT)
print(f"摘要字数: {len(abstract)}")
