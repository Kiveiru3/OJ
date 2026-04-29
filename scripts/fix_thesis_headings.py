import re
from pathlib import Path

from lxml import etree
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\BaiduNetdiskDownload")
LUNWEN_DIR = ROOT / "lunwen1"
INFILE = next(LUNWEN_DIR.glob("基于Spring Boot*.docx"))
OUTFILE = LUNWEN_DIR / "基于Spring Boot的程序设计评测系统的设计与实现-标题格式修订版.docx"


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        paragraph._p = paragraph._element = None


def remove_style_numbering(style):
    ppr = style.element.get_or_add_pPr()
    for node in list(ppr):
        if node.tag == qn("w:numPr"):
            ppr.remove(node)


def ensure_style(doc, name, size_pt=None, font="黑体", bold=None, align=None, before=None, after=None, outline=None):
    styles = doc.styles
    try:
        style = styles[name]
    except Exception:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    if size_pt:
        style.font.size = Pt(size_pt)
    style.font.name = font
    style.font.bold = bold

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), "Times New Roman" if font == "宋体" else font)
    rfonts.set(qn("w:hAnsi"), "Times New Roman" if font == "宋体" else font)

    paragraph_format = style.paragraph_format
    if align is not None:
        paragraph_format.alignment = align
    paragraph_format.left_indent = None
    paragraph_format.first_line_indent = None
    if before is not None:
        paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph_format.space_after = Pt(after)

    remove_style_numbering(style)
    ppr = style.element.get_or_add_pPr()
    if outline is not None:
        for node in list(ppr):
            if node.tag == qn("w:outlineLvl"):
                ppr.remove(node)
        outline_level = OxmlElement("w:outlineLvl")
        outline_level.set(qn("w:val"), str(outline))
        ppr.append(outline_level)

    return style


def set_run_font(run, font="宋体", size=None, bold=None):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), "Times New Roman" if font == "宋体" else font)
    rfonts.set(qn("w:hAnsi"), "Times New Roman" if font == "宋体" else font)


def clear_numbering(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    for node in list(ppr):
        if node.tag == qn("w:numPr"):
            ppr.remove(node)


def set_text(doc, paragraph, text, style=None):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if style is not None:
        paragraph.style = doc.styles[style] if isinstance(style, str) else style
    return paragraph.add_run(text)


def style_paragraph(paragraph):
    if paragraph._element is None:
        return

    clear_numbering(paragraph)
    style_name = paragraph.style.name if paragraph.style is not None else ""

    if style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None
        for run in paragraph.runs:
            set_run_font(run, "黑体", 15, None)
    elif style_name == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None
        for run in paragraph.runs:
            set_run_font(run, "黑体", 14, None)
    elif style_name == "Heading 3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None
        for run in paragraph.runs:
            set_run_font(run, "黑体", 13, True)
    elif style_name in ("Body Text", "Normal"):
        if paragraph.text.strip() and not paragraph.text.strip().startswith(("图 ", "表 ")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(24)
        for run in paragraph.runs:
            set_run_font(run, "宋体", 12, None)


def insert_before(doc, target, text, style="Heading 3"):
    paragraph = target.insert_paragraph_before(text)
    paragraph.style = doc.styles[style]
    style_paragraph(paragraph)
    return paragraph


def find_para(doc, prefix, style_name=None):
    for paragraph in doc.paragraphs:
        if paragraph._element is None:
            continue
        text = " ".join(paragraph.text.strip().split())
        if text.startswith(prefix) and (style_name is None or paragraph.style.name == style_name):
            return paragraph
    return None


def is_toc_title(text):
    compact = text.replace(" ", "").replace("\t", "")
    return compact == "目录"


def add_toc_line(doc, before_paragraph, title, page, level):
    paragraph = before_paragraph.insert_paragraph_before()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = Pt({1: 0, 2: 24, 3: 48}[level])
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Pt(452), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = paragraph.add_run(f"{title}\t{page}")
    set_run_font(run, "黑体" if level == 1 else "宋体", 12, None)
    return paragraph


def attach_toc_section_properties(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    for node in list(ppr):
        if node.tag == qn("w:sectPr"):
            ppr.remove(node)
    sect = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "continuous")
    sect.append(section_type)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:fmt"), "upperRoman")
    sect.append(pg_num)
    pg_size = OxmlElement("w:pgSz")
    pg_size.set(qn("w:w"), "11906")
    pg_size.set(qn("w:h"), "16838")
    sect.append(pg_size)
    pg_mar = OxmlElement("w:pgMar")
    for key, value in {
        "top": "1120",
        "right": "850",
        "bottom": "1333",
        "left": "992",
        "header": "884",
        "footer": "990",
        "gutter": "0",
    }.items():
        pg_mar.set(qn(f"w:{key}"), value)
    sect.append(pg_mar)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:space"), "720")
    sect.append(cols)
    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:linePitch"), "100")
    sect.append(doc_grid)
    ppr.append(sect)


def rebuild_static_toc(doc):
    # The previous draft stores the old TOC in a Word structured-document-tag,
    # which python-docx does not expose as normal paragraphs. Remove it first.
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sdt"):
            xml = etree.tostring(child, encoding="unicode")
            if "Table of Contents" in xml or 'instr="TOC' in xml or ">TOC" in xml:
                body.remove(child)

    toc_title = None
    first_body_heading = None
    for paragraph in doc.paragraphs:
        text = normalized_text(paragraph)
        if toc_title is None and is_toc_title(text):
            toc_title = paragraph
            continue
        if toc_title is not None and text.startswith("第 1 章"):
            first_body_heading = paragraph
            break

    if toc_title is None or first_body_heading is None:
        return

    between = []
    collecting = False
    for paragraph in list(doc.paragraphs):
        if paragraph._element is toc_title._element:
            collecting = True
            continue
        if paragraph._element is first_body_heading._element:
            break
        if collecting:
            between.append(paragraph)
    for paragraph in between:
        delete_paragraph(paragraph)

    toc_entries = [
        (1, "第 1 章 引言", 1),
        (2, "1.1 研究背景", 1),
        (2, "1.2 国内外研究现状", 1),
        (3, "1.2.1 国内研究现状", 1),
        (3, "1.2.2 国外研究现状", 3),
        (2, "1.3 研究内容与目标", 3),
        (2, "1.4 论文组织结构", 3),
        (1, "第 2 章 开发环境与技术简介", 5),
        (2, "2.1 开发环境", 5),
        (2, "2.2 后端关键技术", 5),
        (2, "2.3 前端关键技术", 7),
        (2, "2.4 判题与容器化技术", 7),
        (2, "2.5 本章小结", 7),
        (1, "第 3 章 系统需求分析与总体设计", 9),
        (2, "3.1 可行性分析", 9),
        (3, "3.1.1 技术可行性", 9),
        (3, "3.1.2 经济可行性", 9),
        (3, "3.1.3 操作可行性", 9),
        (2, "3.2 角色需求分析", 9),
        (2, "3.3 系统总体架构设计", 11),
        (2, "3.4 功能结构设计", 11),
        (2, "3.5 关键业务流程设计", 11),
        (3, "3.5.1 代码提交与判题流程", 11),
        (3, "3.5.2 竞赛组织与排名流程", 13),
        (2, "3.6 数据库设计", 14),
        (3, "3.6.1 数据库概念结构设计", 14),
        (3, "3.6.2 核心判题链路设计", 16),
        (3, "3.6.3 数据表结构设计", 16),
        (2, "3.7 本章小结", 18),
        (1, "第 4 章 系统详细设计与实现", 19),
        (2, "4.1 用户与认证模块设计", 19),
        (2, "4.2 题库与测试用例模块设计", 21),
        (2, "4.3 在线评测与判题模块设计", 23),
        (2, "4.4 竞赛管理与排名模块设计", 24),
        (2, "4.5 讨论区与社交模块设计", 26),
        (2, "4.6 教师分析模块设计", 27),
        (2, "4.7 管理控制台设计", 28),
        (2, "4.8 关键实现特性总结", 30),
        (2, "4.9 本章小结", 30),
        (1, "第 5 章 系统测试", 31),
        (2, "5.1 测试环境与测试目标", 31),
        (2, "5.2 功能测试设计", 31),
        (2, "5.3 烟雾测试与端到端验证", 32),
        (2, "5.4 性能测试方案", 33),
        (2, "5.5 本章小结", 33),
        (1, "结论", 34),
        (1, "参考文献", 35),
        (1, "致 谢", 36),
    ]
    last_toc_paragraph = None
    for level, title, page in toc_entries:
        last_toc_paragraph = add_toc_line(doc, first_body_heading, title, page, level)
    if last_toc_paragraph is not None:
        attach_toc_section_properties(last_toc_paragraph)


def normalized_text(paragraph):
    return " ".join(paragraph.text.strip().split())


def replace_paragraph_text(doc, paragraph, revised):
    old_style = paragraph.style
    set_text(doc, paragraph, revised, old_style)
    style_paragraph(paragraph)


def previous_live_paragraph(doc, paragraph):
    previous = None
    for candidate in doc.paragraphs:
        if candidate._element is paragraph._element:
            return previous
        if candidate._element is not None and normalized_text(candidate):
            previous = candidate
    return None


def main():
    doc = Document(str(INFILE))

    ensure_style(doc, "Heading 1", 15, "黑体", None, WD_ALIGN_PARAGRAPH.CENTER, 40, 20, 0)
    ensure_style(doc, "Heading 2", 14, "黑体", None, WD_ALIGN_PARAGRAPH.LEFT, 24, 6, 1)
    ensure_style(doc, "Heading 3", 13, "黑体", True, WD_ALIGN_PARAGRAPH.LEFT, 12, 6, 2)
    try:
        ensure_style(doc, "Body Text", 12, "宋体", None, WD_ALIGN_PARAGRAPH.JUSTIFY, None, 6, None)
    except Exception:
        pass

    replacements = {
        "竞 赛": "竞赛",
        "竞 信息": "竞赛信息",
        "竞 训练": "竞赛训练",
        "赛第一": "第一",
        "赛操作": "操作",
        "赛当前": "当前",
        "赛从工程": "从工程",
        "赛本章": "本章",
        "评分一 致性": "评分一致性",
        "资源统 计 等 步 骤": "资源统计等步骤",
        "、 事中枢": "、赛事中枢",
        " 事中枢": " 赛事中枢",
        "题 目": "题目",
    }
    for paragraph in doc.paragraphs:
        if paragraph._element is None or not paragraph.text:
            continue
        original = paragraph.text
        revised = original
        for old, new in replacements.items():
            revised = revised.replace(old, new)
        revised = re.sub(r"竞\s+赛", "竞赛", revised)
        revised = re.sub(r"竞\s+训练", "竞赛训练", revised)
        revised = re.sub(r"竞\s+测试", "竞赛测试", revised)
        revised = re.sub(r"竞\s+数据", "竞赛数据", revised)
        revised = re.sub(r"竞\s+信息", "竞赛信息", revised)
        revised = re.sub(r"UserController\s+则", "UserController 则", revised)
        revised = re.sub(r"题\s+目", "题目", revised)
        revised = revised.replace("赛上述", "上述")
        revised = revised.replace("赛教师", "教师")
        for lead in ["第一", "操作", "当前", "从工程", "本章"]:
            if revised.startswith("赛" + lead):
                revised = revised[1:]
        if revised != original and len(paragraph._p.xpath(".//w:drawing")) == 0:
            replace_paragraph_text(doc, paragraph, revised)

    heading_renames = {
        "第 1 章 绪论": "第 1 章 引言",
        "第 2 章 相关技术与开发环境": "第 2 章 开发环境与技术简介",
        "第 6 章 结论": "结论",
    }
    for paragraph in doc.paragraphs:
        if paragraph._element is None:
            continue
        text = " ".join(paragraph.text.strip().split())
        if text in heading_renames:
            set_text(doc, paragraph, heading_renames[text], "Heading 1")
            style_paragraph(paragraph)

    # Remove mechanically added single third-level headings under chapter 4 and chapter 5.
    for paragraph in list(doc.paragraphs):
        if paragraph._element is None:
            continue
        text = " ".join(paragraph.text.strip().split())
        if paragraph.style.name == "标题 31" and (text.startswith("4.") or text.startswith("5.")):
            delete_paragraph(paragraph)

    # Keep database subsection, but convert it to a real Heading 3 style.
    for paragraph in doc.paragraphs:
        if paragraph._element is None:
            continue
        text = " ".join(paragraph.text.strip().split())
        if paragraph.style.name == "标题 31" and text.startswith("3.6.1"):
            set_text(doc, paragraph, "3.6.1 数据库概念结构设计", "Heading 3")
            style_paragraph(paragraph)

    # Add third-level headings only where there are two or more real subsections.
    paragraph = find_para(doc, "从已有研究与工程实践来看")
    if paragraph and not any("1.2.1 国内研究现状" in q.text for q in doc.paragraphs if q._element is not None):
        insert_before(doc, paragraph, "1.2.1 国内研究现状")

    paragraph = find_para(doc, "现有研究虽然已经从不同角度探索")
    if paragraph and not any("1.2.2 国外研究现状" in q.text for q in doc.paragraphs if q._element is not None):
        insert_before(doc, paragraph, "1.2.2 国外研究现状")
        extra = paragraph.insert_paragraph_before(
            "国外在线评测平台建设较早，Codeforces、HackerRank、LeetCode、DOMjudge 等平台已在自动判题、竞赛组织、题目标签、排行榜展示和社区讨论方面形成较成熟的实践模式。这些平台通常强调评测效率、并发提交处理、运行环境隔离和竞赛公平性，为高校程序设计训练系统提供了可借鉴的功能形态。"
        )
        extra.style = doc.styles["Body Text"]
        style_paragraph(extra)

    for prefix, title in [
        ("从技术可行性角度看", "3.1.1 技术可行性"),
        ("从经济可行性角度看", "3.1.2 经济可行性"),
        ("从操作可行性角度看", "3.1.3 操作可行性"),
    ]:
        paragraph = find_para(doc, prefix)
        if paragraph and not any(title in q.text for q in doc.paragraphs if q._element is not None):
            insert_before(doc, paragraph, title)

    for prefix, title in [
        ("系统最核心的业务流程是学生提交代码后进入判题", "3.5.1 代码提交与判题流程"),
        ("竞赛模块需要在普通题库练习基础上增加报名", "3.5.2 竞赛组织与排名流程"),
    ]:
        paragraph = find_para(doc, prefix)
        if paragraph and not any(title in q.text for q in doc.paragraphs if q._element is not None):
            insert_before(doc, paragraph, title)

    for prefix, title in [
        ("考虑到论文插图的可读性", "3.6.2 核心判题链路设计"),
        ("在核心表设计中", "3.6.3 数据表结构设计"),
    ]:
        paragraph = find_para(doc, prefix)
        if paragraph and not any(title in q.text for q in doc.paragraphs if q._element is not None):
            insert_before(doc, paragraph, title)

    for paragraph in doc.paragraphs:
        if paragraph._element is None:
            continue
        text = " ".join(paragraph.text.strip().split())
        if text.startswith("全文共分为五章。第 1 章"):
            set_text(
                doc,
                paragraph,
                "全文主体共分为五章，并在正文之后设置结论、参考文献和致谢。第 1 章为引言，说明课题背景、研究现状、研究内容和论文结构。",
                paragraph.style,
            )
            style_paragraph(paragraph)
        elif text.startswith("第 2 章介绍系统所采用的开发环境与关键技术"):
            set_text(
                doc,
                paragraph,
                "第 2 章介绍系统所采用的开发环境与关键技术。第 3 章从角色需求、功能结构、总体架构和数据库设计等方面给出系统总体设计方案。第 4 章是全文重点，对各核心模块的详细设计与实现过程进行阐述。",
                paragraph.style,
            )
            style_paragraph(paragraph)
        elif text.startswith("第 5 章结合现有工程脚本") or text.startswith("细设计与实现过程进行阐述"):
            set_text(
                doc,
                paragraph,
                "第 5 章结合现有工程脚本与测试思路，对系统测试方案进行总结。最后给出结论、参考文献和致谢。",
                paragraph.style,
            )
            style_paragraph(paragraph)

    # Merge or remove broken fragments left by the previous generated draft.
    for paragraph in list(doc.paragraphs):
        if paragraph._element is None:
            continue
        text = normalized_text(paragraph)
        previous = previous_live_paragraph(doc, paragraph)

        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3", "标题 31") and not text:
            delete_paragraph(paragraph)
            continue

        if text.startswith("构和数据库设计四个方面给出系统总体设计方案"):
            delete_paragraph(paragraph)
            continue

        if text == "赛":
            delete_paragraph(paragraph)
            continue

        if previous is not None:
            previous_text = normalized_text(previous)
            if text.startswith("习、代码提交") and previous_text.endswith("题目练"):
                replace_paragraph_text(doc, previous, previous.text + paragraph.text)
                delete_paragraph(paragraph)
                continue
            if text.startswith("互、模块化") and previous_text.endswith("响应式交"):
                replace_paragraph_text(doc, previous, previous.text + paragraph.text)
                delete_paragraph(paragraph)
                continue
            if text.startswith("赛") and previous_text.endswith("竞"):
                replace_paragraph_text(doc, previous, previous.text + paragraph.text)
                delete_paragraph(paragraph)
                continue
            if text.startswith("列表以及管理员接口校验") and previous_text.endswith("讨论"):
                replace_paragraph_text(doc, previous, previous.text + paragraph.text)
                delete_paragraph(paragraph)
                continue

        if text.startswith("量上题和题库迁移过程中"):
            replace_paragraph_text(
                doc,
                paragraph,
                "在批量上题和题库迁移过程中具有较强实用价值。测试用例设计上，系统将样例与隐藏数据分离存储，其中样例可直接展示在题面中，隐藏数据则用于最终判定，有助于提升评测公平性。",
            )

    for paragraph in doc.paragraphs:
        if paragraph._element is None or not paragraph.text or len(paragraph._p.xpath(".//w:drawing")) > 0:
            continue
        revised = paragraph.text
        revised = re.sub(r"竞\s+赛", "竞赛", revised)
        revised = re.sub(r"竞\s+训练", "竞赛训练", revised)
        revised = re.sub(r"竞\s+测试", "竞赛测试", revised)
        revised = re.sub(r"竞\s+数据", "竞赛数据", revised)
        revised = re.sub(r"竞\s+信息", "竞赛信息", revised)
        revised = re.sub(r"UserController\s+则", "UserController 则", revised)
        revised = re.sub(r"题\s+目", "题目", revised)
        revised = revised.replace("赛上述", "上述")
        revised = revised.replace("赛教师", "教师")
        if revised != paragraph.text:
            replace_paragraph_text(doc, paragraph, revised)

    # Remove duplicated empty appendix headings in the previous revision.
    for paragraph in list(doc.paragraphs):
        if paragraph._element is None:
            continue
        text = " ".join(paragraph.text.strip().split())
        if text == "附 录":
            delete_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        if paragraph._element is None:
            continue
        if paragraph.style.name == "标题 31":
            paragraph.style = doc.styles["Heading 3"]
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            clear_numbering(paragraph)
        style_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        text = normalized_text(paragraph)
        if text.startswith("关键词：") or text.startswith("Key words:"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.left_indent = None
            for run in paragraph.runs:
                set_run_font(run, "Times New Roman" if text.startswith("Key words:") else "宋体", 12, None)

    rebuild_static_toc(doc)

    doc.save(str(OUTFILE))
    print(OUTFILE)


if __name__ == "__main__":
    main()
