from __future__ import annotations

import argparse
import json
import re
import struct
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SCREENSHOT_PLACEHOLDER_RE = re.compile(r"^\[此处插入截图：(.+?)\]$")
FIGURE_CAPTION_RE = re.compile(r"^图\s*\d+(\.\d+)?")
TABLE_CAPTION_RE = re.compile(r"^表\s*\d+(\.\d+)?")
INLINE_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
AUTOLINK_RE = re.compile(r"<(https?://[^>]+)>")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a reference-style thesis DOCX from markdown.")
    parser.add_argument("source", type=Path)
    parser.add_argument("reference", type=Path)
    parser.add_argument("target", type=Path)
    parser.add_argument("--image-map", dest="image_map", type=Path, default=None)
    return parser.parse_args()


def cleanup_inline_markdown(text: str) -> str:
    cleaned = text.replace("`", "")
    cleaned = INLINE_LINK_RE.sub(r"\1", cleaned)
    cleaned = AUTOLINK_RE.sub(r"\1", cleaned)
    for pattern, repl in (
        (r"\*\*(.+?)\*\*", r"\1"),
        (r"__(.+?)__", r"\1"),
        (r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"\1"),
        (r"(?<!_)_(?!\s)(.+?)(?<!\s)_(?!_)", r"\1"),
    ):
        cleaned = re.sub(pattern, repl, cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def normalize_marker(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", text)


def compact_heading(text: str) -> str:
    return normalize_marker(text).lower()


def load_image_map(path: Path | None) -> dict[str, Path]:
    if path is None or not path.exists():
        return {}
    raw = json.loads(path.read_text(encoding="utf-8"))
    return {key: Path(value) for key, value in raw.items()}


def set_run_fonts(run, east_asia_font: str, latin_font: str, size_pt: float, *, bold: bool | None = None) -> None:
    if bold is not None:
        run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_paragraph_text(
    paragraph,
    text: str,
    *,
    east_asia_font: str,
    latin_font: str,
    size_pt: float,
    bold: bool | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, east_asia_font, latin_font, size_pt, bold=bold)


def apply_paragraph_format(
    paragraph,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    line_spacing: float | None = None,
    line_spacing_rule: WD_LINE_SPACING = WD_LINE_SPACING.MULTIPLE,
    first_line_indent_pt: float = 0,
    left_indent_pt: float = 0,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing_rule = line_spacing_rule
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    paragraph.paragraph_format.left_indent = Pt(left_indent_pt)
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)


def add_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:pageBreakBefore")) is None:
        page_break_before = OxmlElement("w:pageBreakBefore")
        p_pr.append(page_break_before)


def style_cover_main(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=36)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def style_cover_title(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=26)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def style_cover_meta(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="仿宋", latin_font="Times New Roman", size_pt=16)
    apply_paragraph_format(paragraph, line_spacing=1.5, first_line_indent_pt=495)


def style_cover_date(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Arial", size_pt=16)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)


def style_statement_heading(paragraph) -> None:
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    for run in paragraph.runs:
        set_run_fonts(run, "黑体", "Times New Roman", 15)


def style_statement_body(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="宋体", latin_font="Times New Roman", size_pt=12)
    apply_paragraph_format(paragraph, line_spacing=1.5, first_line_indent_pt=24)


def style_statement_signature(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="宋体", latin_font="Times New Roman", size_pt=12)
    apply_paragraph_format(paragraph, line_spacing=1.5)


def style_abstract_heading(paragraph, english: bool = False) -> None:
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in paragraph.runs:
        set_run_fonts(run, "黑体" if not english else "Times New Roman", "Times New Roman", 15)


def style_body(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="宋体", latin_font="Times New Roman", size_pt=10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.3, first_line_indent_pt=24)


def style_body_en(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="Times New Roman", latin_font="Times New Roman", size_pt=10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.3, first_line_indent_pt=24)


def style_keywords(paragraph, label: str, content: str, *, english: bool = False) -> None:
    clear_paragraph(paragraph)
    label_run = paragraph.add_run(label)
    content_run = paragraph.add_run(content)
    if english:
        set_run_fonts(label_run, "Times New Roman", "Times New Roman", 10.5, bold=True)
        set_run_fonts(content_run, "Times New Roman", "Times New Roman", 10.5)
    else:
        set_run_fonts(label_run, "黑体", "Times New Roman", 10.5, bold=False)
        set_run_fonts(content_run, "宋体", "Times New Roman", 10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def style_heading1(paragraph, text: str) -> None:
    paragraph.style = "Heading 1"
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=15)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def style_heading2(paragraph, text: str) -> None:
    paragraph.style = "Heading 2"
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=14)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def style_heading3(paragraph, text: str) -> None:
    paragraph.style = "Heading 3"
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=13)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def style_figure_caption(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE)


def style_table_caption(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="黑体", latin_font="Times New Roman", size_pt=10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE)


def style_code(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="Consolas", latin_font="Consolas", size_pt=9)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.1, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)


def style_missing_asset(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="楷体", latin_font="Times New Roman", size_pt=10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def style_reference(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, text, east_asia_font="宋体", latin_font="Times New Roman", size_pt=10.5)
    apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=16, line_spacing_rule=WD_LINE_SPACING.EXACTLY, first_line_indent_pt=-21, left_indent_pt=21)


def read_png_size(path: Path) -> tuple[int, int] | None:
    try:
        with path.open("rb") as fh:
            header = fh.read(24)
    except OSError:
        return None
    if len(header) < 24 or header[:8] != b"\x89PNG\r\n\x1a\n" or header[12:16] != b"IHDR":
        return None
    return struct.unpack(">II", header[16:24])


def picture_width_cm(path: Path) -> float:
    size = read_png_size(path)
    if not size:
        return 14.5
    width_px, height_px = size
    ratio = width_px / max(height_px, 1)
    if width_px < 1200:
        return 11.0
    if ratio >= 2.2:
        return 15.2
    if ratio >= 1.5:
        return 14.2
    if ratio >= 1.0:
        return 12.4
    return 10.2


def add_picture(doc: Document, path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(picture_width_cm(path)))


def add_missing_asset(doc: Document, label: str) -> None:
    paragraph = doc.add_paragraph()
    style_missing_asset(paragraph, f"【待补素材：{label}】")


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def add_toc_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    hint_run = OxmlElement("w:r")
    hint_text = OxmlElement("w:t")
    hint_text.text = "目录将在打开文档或导出 PDF 时更新"
    hint_run.append(hint_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run_begin = paragraph.add_run()
    run_begin._r.append(fld_begin)
    run_instr = paragraph.add_run()
    run_instr._r.append(instr)
    run_sep = paragraph.add_run()
    run_sep._r.append(fld_sep)
    paragraph._p.append(hint_run)
    run_end = paragraph.add_run()
    run_end._r.append(fld_end)


def clear_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs:
        clear_paragraph(paragraph)
    for paragraph in section.footer.paragraphs:
        clear_paragraph(paragraph)


def body_children(doc: Document) -> list:
    return list(doc._body._body)


def remove_elements_between(doc: Document, start_paragraph, end_paragraph) -> None:
    body = doc._body._body
    children = body_children(doc)
    start_index = children.index(start_paragraph._element) + 1
    end_index = children.index(end_paragraph._element)
    for child in children[start_index:end_index]:
        body.remove(child)


def remove_from_paragraph_to_end(doc: Document, start_paragraph) -> None:
    body = doc._body._body
    children = body_children(doc)
    start_index = children.index(start_paragraph._element)
    for child in children[start_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def insert_blocks_before(reference_paragraph, blocks: list[tuple[str, str]]) -> None:
    cursor = reference_paragraph
    for block_type, text in reversed(blocks):
        paragraph = cursor.insert_paragraph_before("")
        if block_type == "abstract":
            style_body(paragraph, text)
        elif block_type == "abstract_en":
            style_body_en(paragraph, text)
        elif block_type == "keywords_cn":
            style_keywords(paragraph, "关键词：", text, english=False)
        elif block_type == "keywords_en":
            style_keywords(paragraph, "Key words: ", text, english=True)
        elif block_type == "statement_body":
            style_statement_body(paragraph, text)
        elif block_type == "statement_signature":
            style_statement_signature(paragraph, text)
        else:
            style_body(paragraph, text)
        cursor = paragraph


def find_paragraph(doc: Document, predicate) -> object:
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    raise ValueError("Paragraph not found")


def chinese_date(today: date) -> str:
    digits = {"0": "〇", "1": "一", "2": "二", "3": "三", "4": "四", "5": "五", "6": "六", "7": "七", "8": "八", "9": "九"}
    year = "".join(digits[ch] for ch in str(today.year))
    month = str(today.month)
    day = str(today.day)
    return f"{year}年{month}月{day}日"


def parse_front_matter(lines: list[str]) -> tuple[str, dict[str, str]]:
    title = ""
    meta: dict[str, str] = {}
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            break
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            continue
        if "：" in stripped:
            key, value = stripped.split("：", 1)
            meta[key.strip()] = value.strip()
    return title, meta


def split_sections(lines: list[str]) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    current_heading = ""
    current_lines: list[str] = []
    for line in lines:
        if line.startswith("## "):
            if current_heading:
                sections.append((current_heading, current_lines))
            current_heading = line[3:].strip()
            current_lines = []
        elif current_heading:
            current_lines.append(line.rstrip("\n"))
    if current_heading:
        sections.append((current_heading, current_lines))
    return sections


def section_paragraphs(lines: Iterable[str], *, english: bool = False) -> tuple[list[str], str]:
    paragraphs: list[str] = []
    buffer: list[str] = []
    keywords = ""
    keyword_prefix = "Keywords:" if english else "关键词："
    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            if buffer:
                paragraphs.append(" ".join(buffer))
                buffer = []
            continue
        if stripped.startswith(keyword_prefix):
            if buffer:
                paragraphs.append(" ".join(buffer))
                buffer = []
            keywords = stripped.split("：", 1)[1].strip() if "：" in stripped else stripped.split(":", 1)[1].strip()
            continue
        buffer.append(cleanup_inline_markdown(stripped))
    if buffer:
        paragraphs.append(" ".join(buffer))
    return paragraphs, keywords


def normalize_heading1_text(text: str) -> str:
    compact = compact_heading(text)
    if compact == "致谢":
        return "致 谢"
    if compact == "附录":
        return "附 录"
    if compact == "abstract":
        return "ABSTRACT"
    match = re.match(r"^第\s*(\d+)\s*章\s*(.+)$", text)
    if match:
        return f"第 {match.group(1)} 章 {match.group(2).strip()}"
    return text


def normalize_heading2_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def chapter_number_from_heading(text: str) -> str | None:
    match = re.match(r"^第\s*(\d+)\s*章", text)
    if match:
        return match.group(1)
    return None


def ensure_heading2_number(text: str, chapter_no: str | None, index: int) -> str:
    clean = normalize_heading2_text(text)
    if re.match(r"^\d+\.\d+\s+", clean):
        return clean
    if chapter_no is None:
        return clean
    return f"{chapter_no}.{index} {clean}"


def ensure_heading3_number(text: str, chapter_no: str | None, heading2_index: int, heading3_index: int) -> str:
    clean = normalize_heading2_text(text)
    if re.match(r"^\d+\.\d+\.\d+\s+", clean):
        return clean
    if chapter_no is None or heading2_index <= 0:
        return clean
    return f"{chapter_no}.{heading2_index}.{heading3_index} {clean}"


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in table_lines:
        cells = [cleanup_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r"[:\- ]+", cell or "") for cell in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for candidate in ("Table Grid", "表格网格", "TableNormal"):
        try:
            table.style = candidate
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index, cell_text in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            cell.text = cell_text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_text(paragraph, cell_text, east_asia_font="宋体", latin_font="Times New Roman", size_pt=10.5)
                apply_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE)


def build_body(doc: Document, sections: list[tuple[str, list[str]]], image_map: dict[str, Path]) -> None:
    in_code = False
    code_lang = ""
    pending_mermaid = False
    current_section = ""
    first_heading = True
    current_chapter_no: str | None = None
    heading2_index = 0
    heading3_index = 0

    for heading, section_lines in sections:
        if compact_heading(heading) in {"摘要", "abstract"}:
            continue

        if compact_heading(heading) in {"参考文献", "致谢", "附录"}:
            doc.add_page_break()
            paragraph = doc.add_paragraph()
            style_heading1(paragraph, normalize_heading1_text(heading))
            current_section = compact_heading(heading)
            current_chapter_no = None
            heading2_index = 0
            heading3_index = 0
        elif heading.startswith("第"):
            if not first_heading:
                doc.add_page_break()
            paragraph = doc.add_paragraph()
            heading_text = normalize_heading1_text(heading)
            style_heading1(paragraph, heading_text)
            current_section = compact_heading(heading)
            current_chapter_no = chapter_number_from_heading(heading_text)
            heading2_index = 0
            heading3_index = 0
            first_heading = False
        else:
            paragraph = doc.add_paragraph()
            style_heading1(paragraph, normalize_heading1_text(heading))
            current_section = compact_heading(heading)
            current_chapter_no = None
            heading2_index = 0
            heading3_index = 0

        index = 0
        while index < len(section_lines):
            raw_line = section_lines[index]
            stripped = raw_line.strip()

            if in_code:
                if stripped.startswith("```"):
                    in_code = False
                    if code_lang == "mermaid":
                        pending_mermaid = True
                    code_lang = ""
                    index += 1
                    continue
                if code_lang != "mermaid":
                    paragraph = doc.add_paragraph()
                    style_code(paragraph, raw_line.rstrip())
                index += 1
                continue

            if not stripped:
                index += 1
                continue

            if stripped.startswith("```"):
                in_code = True
                code_lang = stripped[3:].strip().lower()
                index += 1
                continue

            if stripped.startswith("### "):
                paragraph = doc.add_paragraph()
                heading2_index += 1
                heading3_index = 0
                style_heading2(paragraph, ensure_heading2_number(stripped[4:].strip(), current_chapter_no, heading2_index))
                index += 1
                continue

            if stripped.startswith("#### "):
                paragraph = doc.add_paragraph()
                heading3_index += 1
                style_heading3(paragraph, ensure_heading3_number(stripped[5:].strip(), current_chapter_no, heading2_index, heading3_index))
                index += 1
                continue

            screenshot_match = SCREENSHOT_PLACEHOLDER_RE.match(stripped)
            if screenshot_match:
                label = screenshot_match.group(1).strip()
                image_path = image_map.get(label)
                if image_path and image_path.exists():
                    add_picture(doc, image_path)
                else:
                    add_missing_asset(doc, label)
                index += 1
                continue

            if FIGURE_CAPTION_RE.match(stripped):
                if pending_mermaid:
                    image_path = image_map.get(stripped)
                    if image_path and image_path.exists():
                        add_picture(doc, image_path)
                    else:
                        add_missing_asset(doc, stripped)
                    pending_mermaid = False
                paragraph = doc.add_paragraph()
                style_figure_caption(paragraph, cleanup_inline_markdown(stripped))
                index += 1
                continue

            if TABLE_CAPTION_RE.match(stripped):
                paragraph = doc.add_paragraph()
                style_table_caption(paragraph, cleanup_inline_markdown(stripped))
                index += 1
                continue

            if stripped.startswith("|"):
                table_lines: list[str] = []
                while index < len(section_lines) and section_lines[index].strip().startswith("|"):
                    table_lines.append(section_lines[index].strip())
                    index += 1
                add_markdown_table(doc, table_lines)
                continue

            paragraph = doc.add_paragraph()
            clean_text = cleanup_inline_markdown(stripped)
            if current_section == "参考文献" and re.match(r"^\[\d+\]", stripped):
                style_reference(paragraph, clean_text)
            else:
                style_body(paragraph, clean_text)
            index += 1


def update_cover_and_front_pages(doc: Document, title: str, meta: dict[str, str], sections: list[tuple[str, list[str]]]) -> None:
    title_paragraph = find_paragraph(doc, lambda p: "程序设计评测系统的设计与实现" in p.text and normalize_marker(p.text) != "毕业设计（论文）")
    name_paragraph = find_paragraph(doc, lambda p: normalize_marker(p.text).startswith("姓名："))
    student_id_paragraph = find_paragraph(doc, lambda p: normalize_marker(p.text).startswith("学号："))
    department_paragraph = find_paragraph(doc, lambda p: normalize_marker(p.text).startswith("院系："))
    major_paragraph = find_paragraph(doc, lambda p: normalize_marker(p.text).startswith("专业："))
    teacher_paragraph = find_paragraph(doc, lambda p: normalize_marker(p.text).startswith("指导教师："))
    date_paragraph = find_paragraph(doc, lambda p: "二〇" in p.text and "年" in p.text and "月" in p.text)

    def value_or_default(key: str, default: str) -> str:
        value = meta.get(key, "").strip()
        return default if value in {"", "待补充", "待填写", "TBD"} else value

    author = value_or_default("作者", "待补充")
    student_id = value_or_default("学号", "待补充")
    dept_major = value_or_default("院系与专业", "待补充")
    teacher = value_or_default("指导教师", "待补充")
    finish_date = value_or_default("成文日期", chinese_date(date.today()))

    if dept_major in {"", "待补充"}:
        department = "待补充"
        major = "待补充"
    elif any(sep in dept_major for sep in (" / ", "/", "，", ",")):
        for sep in (" / ", "/", "，", ","):
            if sep in dept_major:
                department, major = [part.strip() for part in dept_major.split(sep, 1)]
                break
    else:
        department = dept_major
        major = "待补充"

    style_cover_title(title_paragraph, title)
    style_cover_meta(name_paragraph, f"姓    名：{author}")
    style_cover_meta(student_id_paragraph, f"学    号：{student_id}")
    style_cover_meta(department_paragraph, f"院    系：{department}")
    style_cover_meta(major_paragraph, f"专    业：{major}")
    style_cover_meta(teacher_paragraph, f"指导教师：{teacher}")
    style_cover_date(date_paragraph, finish_date)
    add_page_break_before(date_paragraph)

    originality_heading = find_paragraph(doc, lambda p: normalize_marker(p.text) == "毕业设计（论文）原创性声明")
    authorization_heading = find_paragraph(doc, lambda p: normalize_marker(p.text) == "毕业设计（论文）版权使用授权书")
    abstract_heading = find_paragraph(doc, lambda p: normalize_marker(p.text) == "摘要")
    english_heading = find_paragraph(doc, lambda p: normalize_marker(p.text) == "ABSTRACT")
    toc_heading = find_paragraph(doc, lambda p: normalize_marker(p.text) == "目录")

    statement_date = f"{date.today().year}年 {date.today().month}月 {date.today().day}日"
    statement_blocks = [
        (
            "statement_body",
            f"本人所提交的毕业设计（论文）《{title}》，是在{teacher}导师的指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的研究成果。对本文研究做出重要贡献的个人和集体，均已在文中作出明确说明并表示谢意。",
        ),
        ("statement_body", "本声明的法律后果由本人承担。"),
        ("statement_signature", "论文作者（签名）：            指导教师（签名）："),
        ("statement_signature", f"{statement_date}                        {statement_date}"),
    ]
    remove_elements_between(doc, originality_heading, authorization_heading)
    insert_blocks_before(authorization_heading, statement_blocks)

    authorization_blocks = [
        (
            "statement_body",
            "本人完全了解唐山学院有关保留、使用毕业设计（论文）的规定，同意学校保留并向有关部门或机构送交本论文的复印件和电子版，允许论文被查阅和借阅；同意学校将论文的全部或部分内容编入有关数据库进行检索，并采用影印、缩印或数字化等复制手段保存、汇编本论文。",
        ),
        ("statement_body", "保密的毕业设计（论文）在解密后适用本授权书。"),
        ("statement_signature", "论文作者（签名）：            指导教师（签名）："),
        ("statement_signature", f"{statement_date}                        {statement_date}"),
    ]
    remove_elements_between(doc, authorization_heading, abstract_heading)
    insert_blocks_before(abstract_heading, authorization_blocks)

    for paragraph in (originality_heading, authorization_heading, abstract_heading, english_heading, toc_heading):
        if normalize_marker(paragraph.text) in {"摘要", "ABSTRACT"}:
            style_abstract_heading(paragraph, english=normalize_marker(paragraph.text) == "ABSTRACT")
        else:
            style_statement_heading(paragraph)
    for paragraph in (abstract_heading, english_heading, toc_heading):
        add_page_break_before(paragraph)

    section_map = {heading: lines for heading, lines in sections}
    abstract_paragraphs, abstract_keywords = section_paragraphs(section_map["摘要"], english=False)
    english_paragraphs, english_keywords = section_paragraphs(section_map["Abstract"], english=True)

    remove_elements_between(doc, abstract_heading, english_heading)
    abstract_blocks = [("abstract", text) for text in abstract_paragraphs]
    if abstract_keywords:
        abstract_blocks.append(("keywords_cn", abstract_keywords))
    insert_blocks_before(english_heading, abstract_blocks)

    remove_elements_between(doc, english_heading, toc_heading)
    english_blocks = [("abstract_en", text) for text in english_paragraphs]
    if english_keywords:
        english_blocks.append(("keywords_en", english_keywords))
    insert_blocks_before(toc_heading, english_blocks)

    toc_children = body_children(doc)
    toc_index = toc_children.index(toc_heading._element)
    chapter_start = find_paragraph(doc, lambda p: compact_heading(p.text).startswith("第1章") or compact_heading(p.text).startswith("第1章引言"))
    chapter_start_index = toc_children.index(chapter_start._element)
    if chapter_start_index == toc_index + 1:
        toc_paragraph = chapter_start.insert_paragraph_before("")
        add_toc_field(toc_paragraph)
        style_body(toc_paragraph, "")


def main() -> int:
    args = parse_args()
    lines = args.source.read_text(encoding="utf-8").splitlines()
    title, meta = parse_front_matter(lines)
    sections = split_sections(lines)
    if not title:
        raise ValueError("Markdown source is missing a title line.")

    doc = Document(args.reference)
    update_cover_and_front_pages(doc, title, meta, sections)

    first_chapter = find_paragraph(doc, lambda p: compact_heading(p.text).startswith("第1章") or compact_heading(p.text).startswith("第1章引言"))
    remove_from_paragraph_to_end(doc, first_chapter)

    body_sections = [item for item in sections if compact_heading(item[0]) not in {"摘要", "abstract"}]
    image_map = load_image_map(args.image_map)
    build_body(doc, body_sections, image_map)

    if len(doc.sections) >= 1:
        clear_header_footer(doc.sections[0])
    if len(doc.sections) >= 2:
        clear_header_footer(doc.sections[1])

    section = doc.sections[-1]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    enable_update_fields(doc)
    args.target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.target)
    print(args.target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
